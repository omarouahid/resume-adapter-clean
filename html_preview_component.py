#!/usr/bin/env python3
"""
HTML Preview Component - Live preview for HTML resumes in Streamlit.
"""

import streamlit as st
import streamlit.components.v1 as components
from typing import Dict, Optional, Any
import base64
import logging
import io
import zipfile
import tempfile
import os
from pathlib import Path

logger = logging.getLogger(__name__)

class HTMLResumePreview:
    """Component for displaying live HTML resume previews."""
    
    def __init__(self):
        """Initialize the HTML preview component."""
        self.component_name = "html_resume_preview"
    
    def show_preview(self, 
                    html_content: str, 
                    css_content: str = "",
                    height: int = 800) -> None:
        """
        Display HTML resume preview in Streamlit.
        
        Args:
            html_content: HTML content to display
            css_content: Optional separate CSS content
            height: Height of the preview iframe
        """
        
        if not html_content.strip():
            st.warning("No HTML content to preview")
            return
        
        try:
            # Prepare the complete HTML document
            if css_content:
                # If CSS is provided separately, embed it
                complete_html = self._create_complete_html(html_content, css_content)
            else:
                # Assume HTML already has CSS embedded or inline styles
                complete_html = html_content
            
            # Add preview-specific enhancements
            enhanced_html = self._enhance_html_for_preview(complete_html)
            
            # Display using Streamlit's HTML component
            components.html(
                enhanced_html,
                height=height,
                scrolling=True
            )
            
        except Exception as e:
            st.error(f"Error displaying HTML preview: {str(e)}")
            logger.error(f"HTML preview error: {e}")
    
    def show_side_by_side_preview(self,
                                 original_html: str,
                                 modified_html: str,
                                 original_css: str = "",
                                 modified_css: str = "",
                                 height: int = 700) -> None:
        """
        Show side-by-side comparison of two HTML resumes.
        
        Args:
            original_html: Original HTML content
            modified_html: Modified HTML content
            original_css: CSS for original
            modified_css: CSS for modified
            height: Height of preview iframes
        """
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 📄 Original")
            self.show_preview(
                original_html, 
                original_css, 
                height=height
            )
        
        with col2:
            st.markdown("#### ✨ Enhanced")
            self.show_preview(
                modified_html,
                modified_css,
                height=height
            )
    
    def show_responsive_preview(self,
                              html_content: str,
                              css_content: str = "",
                              show_mobile: bool = True,
                              show_tablet: bool = True,
                              show_desktop: bool = True) -> None:
        """
        Show responsive preview at different screen sizes.
        
        Args:
            html_content: HTML content
            css_content: CSS content
            show_mobile: Show mobile preview
            show_tablet: Show tablet preview
            show_desktop: Show desktop preview
        """
        
        st.markdown("### 📱 Responsive Preview")
        
        # Device selection
        devices = []
        if show_mobile:
            devices.append(("📱 Mobile", 375, 600))
        if show_tablet:
            devices.append(("📱 Tablet", 768, 600))
        if show_desktop:
            devices.append(("💻 Desktop", 1200, 700))
        
        if not devices:
            st.warning("No devices selected for preview")
            return
        
        # Create columns based on number of devices
        cols = st.columns(len(devices))
        
        for i, (device_name, width, height) in enumerate(devices):
            with cols[i]:
                st.markdown(f"#### {device_name}")
                
                # Create device-specific HTML with viewport constraints
                device_html = self._create_device_preview(
                    html_content, 
                    css_content, 
                    width
                )
                
                components.html(
                    device_html,
                    height=height,
                    scrolling=True
                )
    
    def show_template_gallery(self, templates: Dict[str, Any]) -> Optional[str]:
        """
        Show template gallery for selection.
        
        Args:
            templates: Dictionary of template data
            
        Returns:
            Selected template ID or None
        """
        
        st.markdown("### 🎨 Choose Your Template")
        
        if not templates:
            st.warning("No templates available")
            return None
        
        # Create template grid
        cols_per_row = 2
        template_ids = list(templates.keys())
        
        selected_template = None
        
        for i in range(0, len(template_ids), cols_per_row):
            cols = st.columns(cols_per_row)
            
            for j, col in enumerate(cols):
                if i + j < len(template_ids):
                    template_id = template_ids[i + j]
                    template = templates[template_id]
                    
                    with col:
                        # Template preview card
                        with st.container():
                            st.markdown(f"**{template.get('name', template_id)}**")
                            st.markdown(f"*{template.get('category', 'General')}*")
                            st.markdown(template.get('description', ''))
                            
                            if st.button(
                                f"Select {template.get('name', template_id)}", 
                                key=f"select_template_{template_id}",
                                use_container_width=True
                            ):
                                selected_template = template_id
        
        return selected_template
    
    def show_color_palette_selector(self, palettes: Dict[str, Any]) -> Optional[str]:
        """
        Show color palette selector.
        
        Args:
            palettes: Dictionary of color palette data
            
        Returns:
            Selected palette ID or None
        """
        
        st.markdown("### 🎨 Choose Color Scheme")
        
        if not palettes:
            st.warning("No color palettes available")
            return None
        
        # Create palette grid
        cols_per_row = 3
        palette_ids = list(palettes.keys())
        
        selected_palette = None
        
        for i in range(0, len(palette_ids), cols_per_row):
            cols = st.columns(cols_per_row)
            
            for j, col in enumerate(cols):
                if i + j < len(palette_ids):
                    palette_id = palette_ids[i + j]
                    palette = palettes[palette_id]
                    
                    with col:
                        # Color preview card
                        st.markdown(f"**{palette.get('name', palette_id)}**")
                        
                        # Show color swatches
                        primary = palette.get('primary', '#000000')
                        secondary = palette.get('secondary', '#666666')
                        
                        # Create color swatches using HTML
                        color_html = f"""
                        <div style="display: flex; gap: 5px; margin: 10px 0;">
                            <div style="width: 30px; height: 30px; background: {primary}; border-radius: 50%; border: 2px solid #ddd;"></div>
                            <div style="width: 30px; height: 30px; background: {secondary}; border-radius: 50%; border: 2px solid #ddd;"></div>
                        </div>
                        """
                        
                        st.markdown(color_html, unsafe_allow_html=True)
                        
                        if st.button(
                            f"Use {palette.get('name', palette_id)}", 
                            key=f"select_palette_{palette_id}",
                            use_container_width=True
                        ):
                            selected_palette = palette_id
        
        return selected_palette
    
    def _create_complete_html(self, html_content: str, css_content: str) -> str:
        """Create complete HTML document with embedded CSS."""
        
        # Check if HTML already has DOCTYPE and html tags
        if '<!DOCTYPE' in html_content.upper() or '<html' in html_content.lower():
            # HTML is already complete, just add CSS if needed
            if css_content and '<style>' not in html_content and 'style>' not in html_content:
                # Insert CSS into head
                if '</head>' in html_content:
                    html_content = html_content.replace(
                        '</head>',
                        f'<style>\n{css_content}\n</style>\n</head>'
                    )
            return html_content
        
        # Create complete HTML document
        return f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Resume Preview</title>
    <style>
{css_content}
    </style>
</head>
<body>
{html_content}
</body>
</html>
        """
    
    def _enhance_html_for_preview(self, html_content: str) -> str:
        """Add preview-specific enhancements to HTML."""
        
        # Add print styles and preview optimizations
        preview_styles = """
<style>
/* Preview-specific styles */
body {
    margin: 0;
    padding: 20px;
    background: #f5f5f5;
}

/* Print styles */
@media print {
    body {
        background: white;
        margin: 0;
        padding: 0;
    }
    
    .resume-container {
        box-shadow: none !important;
        margin: 0 !important;
    }
}

/* Smooth transitions */
* {
    transition: all 0.2s ease;
}

/* Better focus styles */
*:focus {
    outline: 2px solid #4A90E2;
    outline-offset: 2px;
}
</style>
        """
        
        # Insert preview styles
        if '</head>' in html_content:
            html_content = html_content.replace('</head>', f'{preview_styles}\n</head>')
        elif '<body>' in html_content:
            html_content = html_content.replace('<body>', f'<head>{preview_styles}</head>\n<body>')
        else:
            html_content = preview_styles + html_content
        
        return html_content
    
    def _create_device_preview(self, html_content: str, css_content: str, width: int) -> str:
        """Create device-specific preview with width constraints."""
        
        complete_html = self._create_complete_html(html_content, css_content)
        
        # Add device-specific viewport constraints
        device_styles = f"""
<style>
body {{
    max-width: {width}px;
    margin: 0 auto;
    padding: 10px;
    background: #f0f0f0;
}}

.resume-container {{
    transform-origin: top left;
    max-width: 100%;
}}

/* Mobile optimizations */
@media (max-width: 480px) {{
    body {{
        padding: 5px;
    }}
    
    .resume-container {{
        font-size: 14px;
    }}
}}
</style>
        """
        
        # Add device styles
        if '</head>' in complete_html:
            complete_html = complete_html.replace('</head>', f'{device_styles}\n</head>')
        
        return complete_html
    
    def create_download_links(self, html_content: str, css_content: str = "") -> None:
        """Create enhanced download links with multiple format options."""
        
        st.markdown("### 💾 Download Resume")
        
        # Prepare complete HTML
        complete_html = self._create_complete_html(html_content, css_content)
        
        # Create download format tabs
        tab1, tab2, tab3, tab4 = st.tabs(["📦 ZIP Package", "📄 PDF", "📝 DOCX", "🖼️ Image"])
        
        with tab1:
            st.markdown("#### 📦 HTML + CSS Package")
            st.info("Complete website package with HTML and CSS files")
            
            col1, col2 = st.columns(2)
            with col1:
                if st.button("📦 Create ZIP Package", use_container_width=True, type="primary"):
                    zip_data = self._create_html_css_zip(html_content, css_content)
                    st.download_button(
                        label="📥 Download ZIP",
                        data=zip_data,
                        file_name="resume_package.zip",
                        mime="application/zip",
                        use_container_width=True
                    )
                    st.success("✅ ZIP package ready!")
            
            with col2:
                # Individual file downloads
                html_bytes = complete_html.encode('utf-8')
                st.download_button(
                    label="📄 Download HTML",
                    data=html_bytes,
                    file_name="resume.html",
                    mime="text/html",
                    use_container_width=True
                )
                
                if css_content:
                    css_bytes = css_content.encode('utf-8')
                    st.download_button(
                        label="🎨 Download CSS",
                        data=css_bytes,
                        file_name="resume.css",
                        mime="text/css",
                        use_container_width=True
                    )
        
        with tab2:
            st.markdown("#### 📄 PDF Export")
            self._create_pdf_download_section(complete_html)
        
        with tab3:
            st.markdown("#### 📝 DOCX Export") 
            self._create_docx_download_section(html_content, css_content)
        
        with tab4:
            st.markdown("#### 🖼️ Image Export")
            self._create_image_download_section(complete_html)
    
    def _create_html_css_zip(self, html_content: str, css_content: str = "") -> bytes:
        """Create a ZIP file containing HTML and CSS files."""
        try:
            # Create ZIP in memory
            zip_buffer = io.BytesIO()
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                # Add complete HTML file
                complete_html = self._create_complete_html(html_content, css_content)
                zip_file.writestr("resume.html", complete_html)
                
                # Add separate CSS file if provided
                if css_content:
                    zip_file.writestr("resume.css", css_content)
                
                # Add a README file
                readme_content = """# Resume Package

## Files Included:
- resume.html: Complete resume with embedded or linked CSS
- resume.css: Separate CSS file (if applicable)

## How to Use:
1. Open resume.html in any web browser
2. For printing/PDF: Press Ctrl+P and save as PDF
3. For customization: Edit the CSS file and refresh the HTML

## Browser Compatibility:
- Chrome, Firefox, Safari, Edge
- Mobile browsers supported

Generated by Resume Adapter
"""
                zip_file.writestr("README.txt", readme_content)
                
                # Add print-optimized version
                print_html = self._create_print_optimized_html(html_content, css_content)
                zip_file.writestr("resume_print.html", print_html)
            
            zip_buffer.seek(0)
            return zip_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error creating ZIP package: {str(e)}")
            raise
    
    def _create_print_optimized_html(self, html_content: str, css_content: str = "") -> str:
        """Create a print-optimized version of the HTML."""
        complete_html = self._create_complete_html(html_content, css_content)
        
        # Add print-specific CSS
        print_css = """
        <style>
        /* Print-specific optimizations */
        @media print {
            * {
                -webkit-print-color-adjust: exact !important;
                color-adjust: exact !important;
                print-color-adjust: exact !important;
            }
            body { 
                background: white !important; 
                color: black !important;
                font-size: 11pt !important;
                line-height: 1.4 !important;
                margin: 0 !important;
                padding: 0 !important;
            }
            .resume-container, .container, .main { 
                box-shadow: none !important;
                max-width: none !important;
                margin: 0 !important;
                padding: 0 !important;
                border: none !important;
            }
            @page {
                margin: 0.5in;
                size: A4;
            }
            .no-print {
                display: none !important;
            }
            a {
                color: black !important;
                text-decoration: none !important;
            }
            .page-break {
                page-break-before: always;
            }
        }
        </style>
        """
        
        # Insert print styles before closing head tag
        if '</head>' in complete_html:
            complete_html = complete_html.replace('</head>', f'{print_css}\n</head>')
        else:
            # If no head tag, add it
            complete_html = f"<head>{print_css}</head>\n{complete_html}"
        
        return complete_html
    
    def _create_pdf_download_section(self, complete_html: str) -> None:
        """Create PDF download options."""
        col1, col2 = st.columns(2)
        
        with col1:
            st.info("🖨️ **Browser Method (Recommended)**")
            st.markdown("""
            1. Download HTML file
            2. Open in browser
            3. Press `Ctrl+P` (Windows) or `Cmd+P` (Mac)
            4. Select "Save as PDF"
            5. Choose A4 paper size
            """)
            
            # Print-optimized HTML download
            print_html = self._create_print_optimized_html(complete_html, "")
            st.download_button(
                label="📄 Download Print-Ready HTML",
                data=print_html.encode('utf-8'),
                file_name="resume_print_ready.html",
                mime="text/html",
                use_container_width=True
            )
        
        with col2:
            st.info("🔧 **Advanced PDF Tools**")
            
            # Try to use wkhtmltopdf if available
            if st.button("🔄 Generate PDF (Server)", use_container_width=True):
                try:
                    pdf_data = self._html_to_pdf_server(complete_html)
                    if pdf_data:
                        st.download_button(
                            label="📥 Download PDF",
                            data=pdf_data,
                            file_name="resume.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                        st.success("✅ PDF generated successfully!")
                    else:
                        st.error("❌ PDF generation failed. Please use browser method.")
                except Exception as e:
                    st.error(f"❌ PDF generation error: {str(e)}")
                    st.info("💡 Please use the browser method instead.")
    
    def _html_to_pdf_server(self, html_content: str) -> Optional[bytes]:
        """Convert HTML to PDF using server-side tools."""
        try:
            # Try using weasyprint if available
            try:
                import weasyprint
                from io import BytesIO
                
                pdf_buffer = BytesIO()
                weasyprint.HTML(string=html_content).write_pdf(pdf_buffer)
                pdf_buffer.seek(0)
                return pdf_buffer.getvalue()
                
            except ImportError:
                pass
            
            # Try using pdfkit if available
            try:
                import pdfkit
                
                options = {
                    'page-size': 'A4',
                    'margin-top': '0.5in',
                    'margin-right': '0.5in', 
                    'margin-bottom': '0.5in',
                    'margin-left': '0.5in',
                    'encoding': 'UTF-8',
                    'no-outline': None,
                    'enable-local-file-access': None
                }
                
                pdf_data = pdfkit.from_string(html_content, False, options=options)
                return pdf_data
                
            except (ImportError, OSError):
                pass
                
            return None
            
        except Exception as e:
            logger.error(f"PDF generation failed: {str(e)}")
            return None
    
    def _create_docx_download_section(self, html_content: str, css_content: str = "") -> None:
        """Create DOCX download options."""
        col1, col2 = st.columns(2)
        
        with col1:
            st.info("📝 **Word Document Export**")
            st.markdown("""
            Convert your resume to Microsoft Word format for easy editing and ATS compatibility.
            """)
            
            if st.button("📝 Generate DOCX", use_container_width=True, type="primary"):
                try:
                    docx_data = self._html_to_docx(html_content)
                    if docx_data:
                        st.download_button(
                            label="📥 Download DOCX",
                            data=docx_data,
                            file_name="resume.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                        st.success("✅ DOCX file generated!")
                    else:
                        st.error("❌ DOCX generation failed")
                        
                except Exception as e:
                    st.error(f"❌ Error generating DOCX: {str(e)}")
        
        with col2:
            st.info("💡 **DOCX Benefits**")
            st.markdown("""
            - **ATS Friendly**: Machine-readable format
            - **Editable**: Easy to customize
            - **Universal**: Opens in Word, Google Docs
            - **Professional**: Standard business format
            """)
    
    def _html_to_docx(self, html_content: str) -> Optional[bytes]:
        """Convert HTML content to DOCX format."""
        try:
            # Try using python-docx with html parsing
            try:
                from docx import Document
                from bs4 import BeautifulSoup
                import re
                
                # Parse HTML
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Create new document
                doc = Document()
                
                # Extract and add content
                self._add_html_to_docx(soup, doc)
                
                # Save to bytes
                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                return docx_buffer.getvalue()
                
            except ImportError:
                st.warning("💡 Install 'python-docx' and 'beautifulsoup4' for DOCX export")
                return None
                
        except Exception as e:
            logger.error(f"DOCX generation failed: {str(e)}")
            return None
    
    def _add_html_to_docx(self, soup, doc):
        """Add HTML content to DOCX document."""
        # Extract text content and structure
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol', 'div']):
            text = element.get_text().strip()
            if not text:
                continue
                
            if element.name in ['h1', 'h2', 'h3', 'h4']:
                # Add as heading
                level = int(element.name[1])
                doc.add_heading(text, level=min(level, 3))
            elif element.name in ['ul', 'ol']:
                # Add list items
                for li in element.find_all('li'):
                    li_text = li.get_text().strip()
                    if li_text:
                        p = doc.add_paragraph(li_text, style='List Bullet' if element.name == 'ul' else 'List Number')
            elif element.name == 'p':
                # Add paragraph
                doc.add_paragraph(text)
            elif element.name == 'div' and text:
                # Add div content as paragraph if it has direct text
                if not element.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol']):
                    doc.add_paragraph(text)
    
    def _create_image_download_section(self, complete_html: str) -> None:
        """Create image download options."""
        col1, col2 = st.columns(2)
        
        with col1:
            st.info("🖼️ **Image Export**")
            st.markdown("""
            Convert your resume to PNG or JPG image format for social media, portfolios, or quick previews.
            """)
            
            # Image format selection
            img_format = st.selectbox("Format:", ["PNG (Transparent)", "JPG (White Background)"])
            
            if st.button("🖼️ Generate Image", use_container_width=True, type="primary"):
                try:
                    img_data = self._html_to_image(complete_html, img_format)
                    if img_data:
                        file_ext = "png" if "PNG" in img_format else "jpg"
                        mime_type = f"image/{file_ext}"
                        
                        st.download_button(
                            label=f"📥 Download {file_ext.upper()}",
                            data=img_data,
                            file_name=f"resume.{file_ext}",
                            mime=mime_type,
                            use_container_width=True
                        )
                        st.success(f"✅ {file_ext.upper()} image generated!")
                    else:
                        st.error("❌ Image generation failed")
                        
                except Exception as e:
                    st.error(f"❌ Error generating image: {str(e)}")
        
        with col2:
            st.info("📱 **Image Uses**")
            st.markdown("""
            - **LinkedIn**: Profile background
            - **Portfolio**: Quick preview
            - **Email**: Inline attachment
            - **Social Media**: Sharing
            """)
    
    def _html_to_image(self, html_content: str, format_type: str) -> Optional[bytes]:
        """Convert HTML to image using available tools."""
        try:
            # Try using the existing image converter
            try:
                from image_converter import image_converter
                
                # Use existing image conversion system
                with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False) as f:
                    f.write(html_content)
                    html_path = f.name
                
                try:
                    # Set format 
                    img_format = "png" if "PNG" in format_type else "jpeg"
                    
                    # Convert using existing system
                    img_data = image_converter.convert_html_file(html_path, format=img_format)
                    
                    if img_data:
                        return img_data
                    else:
                        # Fallback: use screenshot functionality if available
                        return self._html_to_image_fallback(html_content, format_type)
                        
                finally:
                    # Clean up temp file
                    try:
                        os.unlink(html_path)
                    except:
                        pass
                        
            except ImportError:
                return self._html_to_image_fallback(html_content, format_type)
                
        except Exception as e:
            logger.error(f"Image generation failed: {str(e)}")
            return None
    
    def _html_to_image_fallback(self, html_content: str, format_type: str) -> Optional[bytes]:
        """Fallback image generation method."""
        try:
            # Try using selenium if available
            try:
                from selenium import webdriver
                from selenium.webdriver.chrome.options import Options
                import base64
                
                # Setup Chrome options
                chrome_options = Options()
                chrome_options.add_argument('--headless')
                chrome_options.add_argument('--no-sandbox')
                chrome_options.add_argument('--disable-dev-shm-usage')
                chrome_options.add_argument('--window-size=1200,1600')
                
                # Create temporary HTML file
                with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False) as f:
                    f.write(html_content)
                    html_path = f.name
                
                try:
                    driver = webdriver.Chrome(options=chrome_options)
                    driver.get(f"file://{html_path}")
                    
                    # Take screenshot
                    screenshot = driver.get_screenshot_as_png()
                    driver.quit()
                    
                    return screenshot
                    
                finally:
                    try:
                        os.unlink(html_path)
                    except:
                        pass
                        
            except ImportError:
                st.info("💡 Install 'selenium' and Chrome driver for image export")
                return None
                
        except Exception as e:
            logger.error(f"Fallback image generation failed: {str(e)}")
            return None

# Global instance for easy import
html_preview = HTMLResumePreview()